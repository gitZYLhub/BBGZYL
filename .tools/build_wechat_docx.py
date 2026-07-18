from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Civilization\Civ6mods\BBGZYL")
SOURCE = ROOT / "ZYL_LightweightBalance" / "docs" / "full_changes_mobile.md"
OUTPUT = ROOT / "ZYL_LightweightBalance" / "docs" / "ZYL轻量化平衡模组_修改大全_微信版_无编号版.docx"

FONT = "Microsoft YaHei"
BODY = RGBColor(35, 38, 42)
HEADING = RGBColor(31, 77, 120)
HEADING_LIGHT = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 99, 107)
CALLOUT_FILL = "EAF2F8"

PURE_UNCHANGED_MARKERS = (
    "不变",
    "保持原版",
    "不受影响",
    "不增加",
    "不影响",
    "仍需",
    "仍通过",
    "保留四个结社各自的原版发现条件",
    "其他伟人仍按对应区域类型提供产出",
)


def set_style_font(style, size, color=BODY, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def set_cell_free_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep(paragraph, with_next=False, together=False):
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.widow_control = True


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([begin, instruction, separate, display, end])

    run = paragraph.add_run(" 页")
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\[[^\]]+\]\([^\)]+\))")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
            run.font.size = Pt(11)
            run.font.color.rgb = HEADING
        elif part.startswith("["):
            match = re.match(r"\[([^\]]+)\]\([^\)]+\)", part)
            run = paragraph.add_run(match.group(1) if match else part)
        else:
            paragraph.add_run(part)


def add_decimal_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level_specs = [
        (0, "%1.", 500, 260),
        (1, "%1.%2", 760, 340),
    ]
    for level_index, text, left, hanging in level_specs:
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(level_index))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        level.extend([start, number_format, level_text, suffix])

        if level_index == 1:
            restart = OxmlElement("w:lvlRestart")
            restart.set(qn("w:val"), "1")
            level.append(restart)

        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(left))
        indent.set(qn("w:hanging"), str(hanging))
        ppr.extend([tabs, indent])
        level.append(ppr)

        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT)
        rpr.append(rfonts)
        level.append(rpr)
        abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def new_numbering_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def apply_numbering(paragraph, num_id, level_index=0):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), str(level_index))
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numpr.extend([level, number])
    ppr.append(numpr)


def read_and_validate_source():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    unchanged_lines = [
        line.strip()
        for line in lines
        if any(marker in line for marker in PURE_UNCHANGED_MARKERS)
    ]
    if unchanged_lines:
        details = "\n".join(f"- {line}" for line in unchanged_lines)
        raise ValueError(f"手机版只记录实际修改，请删除未修改项：\n{details}")
    return lines


def extract_version(lines):
    for line in lines:
        match = re.fullmatch(r"\*\*当前版本：([^*]+)\*\*", line.strip())
        if match:
            return match.group(1).strip()
    raise ValueError("未在手机版源文档中找到当前版本号")


def ordered_mobile_lines(lines):
    sections = {}
    current = None
    for line in lines:
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = [line]
        elif current is not None:
            sections[current].append(line)

    directory = [
        "## 目录",
        "",
        "- 城市与经济",
        "- 万神殿",
        "- 秘密结社",
        "- 资源与改良设施",
        "- 专家",
        "- 文明调整",
        "- 总督",
        "- 兼容性与实测",
        "",
    ]
    order = [
        "城市与经济",
        "万神殿",
        "秘密结社",
        "资源与改良设施",
        "专家",
        "文明调整",
        "总督",
        "兼容性与实测",
    ]
    expected_sections = set(order) | {"快速目录"}
    missing_sections = set(order) - set(sections)
    unexpected_sections = set(sections) - expected_sections
    if missing_sections or unexpected_sections:
        raise ValueError(
            "手机版章节与微信版生成顺序不一致："
            f"缺少 {sorted(missing_sections)}；未纳入 {sorted(unexpected_sections)}"
        )

    output = list(directory)
    for heading in order:
        output.extend(sections[heading])
    return output


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, 12.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    set_style_font(title, 25, HEADING, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.line_spacing = 1.0

    h1 = styles["Heading 1"]
    set_style_font(h1, 18, HEADING, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.1

    h2 = styles["Heading 2"]
    set_style_font(h2, 15, HEADING_LIGHT, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.1

    h3 = styles["Heading 3"]
    set_style_font(h3, 13, BODY, True)
    h3.paragraph_format.space_before = Pt(9)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.15

    if "Mobile Directory" not in styles:
        numbered = styles.add_style("Mobile Directory", WD_STYLE_TYPE.PARAGRAPH)
    else:
        numbered = styles["Mobile Directory"]
    set_style_font(numbered, 12.5, HEADING, True)
    numbered.paragraph_format.left_indent = Inches(0.12)
    numbered.paragraph_format.space_before = Pt(0)
    numbered.paragraph_format.space_after = Pt(5)
    numbered.paragraph_format.line_spacing = 1.25

    if "Mobile Detail" not in styles:
        detail = styles.add_style("Mobile Detail", WD_STYLE_TYPE.PARAGRAPH)
    else:
        detail = styles["Mobile Detail"]
    set_style_font(detail, 12.5)
    detail.paragraph_format.left_indent = Inches(0.22)
    detail.paragraph_format.right_indent = Inches(0.02)
    detail.paragraph_format.space_before = Pt(0)
    detail.paragraph_format.space_after = Pt(4)
    detail.paragraph_format.line_spacing = 1.25

    if "Mobile Meta" not in styles:
        meta = styles.add_style("Mobile Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Mobile Meta"]
    set_style_font(meta, 11.5, MUTED, False)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.line_spacing = 1.15

    if "Mobile Callout" not in styles:
        callout = styles.add_style("Mobile Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Mobile Callout"]
    set_style_font(callout, 12, HEADING, True)
    callout.paragraph_format.left_indent = Inches(0.10)
    callout.paragraph_format.right_indent = Inches(0.10)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.2


def build_document():
    source_lines = read_and_validate_source()
    version = extract_version(source_lines)

    doc = Document()
    section = doc.sections[0]

    # Mobile-page override to compact_reference_guide: 9:16 portrait canvas.
    section.page_width = Inches(6.0)
    section.page_height = Inches(10.67)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    configure_styles(doc)

    doc.core_properties.title = "ZYL的轻量化平衡模组：修改大全（微信阅读版）"
    doc.core_properties.subject = f"文明6轻量化平衡模组 {version} 完整修改说明"
    doc.core_properties.author = "ZYL"
    doc.core_properties.keywords = "文明6, 模组, 平衡, 微信阅读版"

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("ZYL轻量化平衡模组 · 修改大全")
    hr.font.name = FONT
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)

    title = doc.add_paragraph(style="Title")
    title.add_run("ZYL的轻量化平衡模组")
    set_keep(title, with_next=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("修改大全 · 微信阅读版")
    run.font.name = FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = HEADING_LIGHT
    set_keep(subtitle, with_next=True)

    meta = doc.add_paragraph(style="Mobile Meta")
    meta.add_run(f"当前版本：{version}")
    set_keep(meta, with_next=True)

    meta = doc.add_paragraph(style="Mobile Meta")
    meta.add_run("比较对象：文明 6 原版规则")

    callout = doc.add_paragraph(style="Mobile Callout")
    callout.add_run("适合直接发送到微信。采用窄页、大字号和短段落，打开后即可连续滚动阅读。")
    set_cell_free_paragraph_shading(callout, CALLOUT_FILL)
    set_keep(callout, together=True)

    lines = ordered_mobile_lines(source_lines)
    in_directory = False

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("## "):
            heading = line[3:]
            in_directory = heading == "目录"
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, heading)
            set_keep(p, with_next=True)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
            set_keep(p, with_next=True)
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:])
            set_keep(p, with_next=True)
        elif line.startswith("- "):
            style = "Mobile Directory" if in_directory else "Mobile Detail"
            p = doc.add_paragraph(style=style)
            add_inline(p, line[2:])
            set_keep(p, together=True)
        elif line.startswith("**") and line.endswith("**"):
            label = line[2:-2]
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(label)
            set_keep(p, with_next=True)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
            set_keep(p, together=True)

    doc.save(OUTPUT)
    print(f"WROTE {OUTPUT}")
    print(f"PARAGRAPHS {len(doc.paragraphs)}")


if __name__ == "__main__":
    build_document()
